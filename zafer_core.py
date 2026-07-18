"""
zafer_core.py — Ders materyalleri üzerinde RAG tabanlı asistan çekirdegi.

Sorumluluklar:
  - Materyal klasorunu tarayip parcalara (chunk) bolerek ChromaDB'ye yazmak
  - Degismemis dosyalari yeniden islemeden atlamak (hash tabanli)
  - Soru geldiginde ilgili parcalari bulup LLM'e context olarak vermek
  - Konusma gecmisini sinirli tutmak
"""

from __future__ import annotations

import hashlib
import json
import os
import pathlib
from dataclasses import dataclass, field

import chromadb
from chromadb.utils import embedding_functions
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

# --------------------------------------------------------------------------
# Ayarlar
# --------------------------------------------------------------------------

DB_PATH = os.getenv("ZAFER_DB_PATH", "./zafer_hoca_db")
COLLECTION = "zafer_hoca"

# Turkce icin cok dilli embedding modeli. Varsayilan all-MiniLM-L6-v2
# Ingilizce odaklidir ve Turkce materyallerde arama kalitesi cok dusuktur.
EMBED_MODEL = os.getenv("ZAFER_EMBED_MODEL", "intfloat/multilingual-e5-base")

# OpenRouter model kimligi. Gecerli bir kimlik girmelisin —
# guncel listeyi https://openrouter.ai/models adresinden kontrol et.
LLM_MODEL = os.getenv("ZAFER_LLM_MODEL", "google/gemini-2.5-flash")

CHUNK_CHARS = 1500       # her parcanin yaklasik uzunlugu
CHUNK_OVERLAP = 200      # parcalar arasi bindirme (baglam kopmasin diye)
N_RESULTS = 6            # her soruda cekilecek parca sayisi
MAX_HISTORY_TURNS = 8    # hafizada tutulacak soru-cevap cifti sayisi

SUPPORTED = {".ipynb", ".py", ".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx"}


# --------------------------------------------------------------------------
# Dosya okuma
# --------------------------------------------------------------------------

def read_ipynb(path: pathlib.Path) -> str:
    """Notebook'u markdown + kod hucreleri olarak duz metne cevirir."""
    nb = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    parts = []
    for cell in nb.get("cells", []):
        source = "".join(cell.get("source", [])).strip()
        if not source:
            continue
        if cell.get("cell_type") == "markdown":
            parts.append(f"### ACIKLAMA:\n{source}")
        elif cell.get("cell_type") == "code":
            parts.append(f"### KOD:\n{source}")
    return "\n\n".join(parts)


def read_plain(path: pathlib.Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def read_pdf(path: pathlib.Path) -> str:
    """PDF'ten metin cikarir. Taranmis (resim) PDF'lerde metin katmani yoktur."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(p for p in pages if p)
    if not text:
        print(f"  [!] {path.name}: metin katmani yok (taranmis PDF olabilir).")
    return text


def read_docx(path: pathlib.Path) -> str:
    """Word belgesinden paragraflari ve tablolari cikarir."""
    import docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path: pathlib.Path, max_rows: int = 300) -> str:
    """Excel sayfalarini TSV benzeri metne cevirir (sayfa basi max_rows satir)."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## Sayfa: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                parts.append(f"... ({ws.title}: ilk {max_rows} satir alindi)")
                break
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def read_any(path: pathlib.Path) -> str:
    """Desteklenen her uzanti icin metin dondurur. Hata durumunda neden yazilir."""
    try:
        if path.suffix == ".ipynb":
            return read_ipynb(path)
        if path.suffix == ".pdf":
            return read_pdf(path)
        if path.suffix == ".docx":
            return read_docx(path)
        if path.suffix == ".xlsx":
            return read_xlsx(path)
        if path.suffix in SUPPORTED:
            return read_plain(path)
    except Exception as exc:                      # ciplak except yerine: nedeni gorelim
        print(f"  [!] {path.name} okunamadi: {type(exc).__name__}: {exc}")
    return ""


# --------------------------------------------------------------------------
# Parcalama
# --------------------------------------------------------------------------

def chunk_text(text: str, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP):
    """
    Metni bindirmeli parcalara boler. Mumkun oldugunca satir sonunda keser,
    boylece kod bloklarinin ortasindan bolunme ihtimali azalir.
    """
    if len(text) <= size:
        return [text]

    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            newline = text.rfind("\n", start + size // 2, end)
            if newline != -1:
                end = newline
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


# --------------------------------------------------------------------------
# Vektor deposu
# --------------------------------------------------------------------------

def get_collection():
    embedder = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=EMBED_MODEL
    )
    client = chromadb.PersistentClient(path=DB_PATH)
    return client.get_or_create_collection(
        name=COLLECTION,
        embedding_function=embedder,
        metadata={"hnsw:space": "cosine"},
    )


def index_folder(folder: str | pathlib.Path, force: bool = False) -> int:
    """
    Klasoru tarar, degismis/yeni dosyalari parcalayip veritabanina yazar.
    Ayni dosya tekrar islenirse upsert edilir; ID cakismasi olmaz.
    Dondurulen deger: bu calistirmada islenen dosya sayisi.
    """
    folder = pathlib.Path(folder)
    if not folder.is_dir():
        raise FileNotFoundError(f"Klasor bulunamadi: {folder}")

    col = get_collection()
    processed = 0

    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.suffix not in SUPPORTED:
            continue
        # Jupyter'in otomatik yedek klasorlerini atla
        if ".ipynb_checkpoints" in path.parts:
            continue

        text = read_any(path)
        if len(text) < 50:
            continue

        rel = str(path.relative_to(folder))
        # Dosya yolunu icerige ekle: klasor/dosya adlari da aranabilir olsun.
        # "TIME SERIES/TimeSeriesAnalysis.ipynb" gibi yollar, icerigi kisa
        # dosyalarin bile konu aramasinda bulunmasini saglar.
        text = f"[KAYNAK: {rel}]\n{text}"
        file_key = hashlib.sha1(rel.encode()).hexdigest()[:12]
        content_hash = hashlib.sha1(text.encode()).hexdigest()[:12]

        if not force:
            existing = col.get(where={"file_key": file_key}, limit=1)
            metas = existing.get("metadatas") or []
            if metas and metas[0].get("content_hash") == content_hash:
                continue  # dosya degismemis, atla

        # Eski parcalari temizle (dosya kisalmis olabilir)
        col.delete(where={"file_key": file_key})

        chunks = chunk_text(text)
        col.upsert(
            ids=[f"{file_key}_{i}" for i in range(len(chunks))],
            documents=chunks,
            metadatas=[
                {
                    "src": rel,
                    "type": path.suffix.lstrip("."),
                    "file_key": file_key,
                    "content_hash": content_hash,
                    "chunk": i,
                }
                for i in range(len(chunks))
            ],
        )
        processed += 1
        print(f"  [+] {rel} -> {len(chunks)} parca")

    print(f"\nToplam {processed} dosya islendi. Veritabaninda {col.count()} parca var.")
    return processed


# --------------------------------------------------------------------------
# Ajan
# --------------------------------------------------------------------------

SYSTEM_TEMPLATE = """Sen Zafer Hoca'nin ders materyallerine dayanan bir veri bilimi asistanisin.

Calisma bicimin:
1. Once problemin turunu belirle ve HANGI teknigi neden sectigini acikla.
2. Sectigin teknigin zayif yonlerini de soyle — abartili vaatte bulunma.
3. Calisan, eksiksiz Python kodu yaz. Kod Ingilizce, aciklama Turkce olsun.
4. Asagidaki ders materyallerinde ilgili bir yaklasim varsa ONA oncelik ver
   ve hangi kaynaktan geldigini belirt.
5. Materyallerde cevap yoksa bunu acikca soyle, uydurma.
6. Kullanici bir odev/proje dosyasi (PDF, Word vb.) yuklediginde:
   once problemin ne istedigini 2-3 cumleyle ozetle, sonra bu problemin
   hangi ders materyalindeki hangi teknikle cozulecegini soyle
   (ornek format: "Bu proje, X dosyasinda islenen ARIMA yaklasimiyla
   cozulebilir; su adimlarla ilerle: ..."). Dosya adi vererek yonlendir.

Yon gosterici baslangic noktalari (materyaller aksini soylerse onlari izle):
- Zaman serisi/finans -> once klasik baseline (ARIMA, gradient boosting),
  sonra gerekiyorsa LSTM/GRU. Ham fiyatla LSTM genelde naif tahmin uretir.
- Siniflandirma -> XGBoost / LightGBM / Random Forest
- Kumeleme -> K-Means, DBSCAN, Hierarchical
- Goruntu -> CNN, Transfer Learning
- NLP -> Transformer tabanli modeller
- Oneri sistemi -> Collaborative Filtering, RAG

## Ders Materyalleri:
{context}
"""


@dataclass
class ZaferHoca:
    """Konusma durumunu tasiyan asistan. Her proje icin ayri nesne kullanabilirsin."""

    model: str = LLM_MODEL
    history: list = field(default_factory=list)
    _col: object = None
    _client: object = None

    def __post_init__(self):
        api_key = os.getenv("OPENROUTER_API_KEY")
        if not api_key:
            raise RuntimeError(
                "OPENROUTER_API_KEY tanimli degil. .env dosyasina ekle:\n"
                "  OPENROUTER_API_KEY=sk-or-v1-..."
            )
        self._client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
        self._col = get_collection()

    # -- yardimcilar --------------------------------------------------------

    def retrieve(self, query: str, n: int = N_RESULTS) -> str:
        """Soruyla ilgili materyal parcalarini tek bir context metnine cevirir."""
        if self._col.count() == 0:
            return "(Veritabani bos — once index_folder() calistir.)"

        res = self._col.query(query_texts=[query], n_results=min(n, self._col.count()))
        docs = res["documents"][0]
        metas = res["metadatas"][0]
        dists = res.get("distances", [[None] * len(docs)])[0]

        parts = []
        for doc, meta, dist in zip(docs, metas, dists):
            score = f" (benzerlik: {1 - dist:.2f})" if dist is not None else ""
            parts.append(f"--- Kaynak: {meta['src']}{score} ---\n{doc}")
        return "\n\n".join(parts)

    def reset(self):
        """Yeni projeye baslarken konusma gecmisini temizler."""
        self.history = []

    def _trim(self):
        """Gecmisi son N tur ile sinirlar; context patlamasini onler."""
        keep = MAX_HISTORY_TURNS * 2
        if len(self.history) > keep:
            self.history = self.history[-keep:]

    # -- ana giris noktasi --------------------------------------------------

    def ask(self, message: str, attached_text: str | None = None) -> str:
        """
        Soru sorar. attached_text verilirse (yuklenen dosyanin icerigi)
        soruya ek baglam olarak eklenir.
        """
        user_content = message
        if attached_text:
            user_content = (
                f"{message}\n\n## Ekteki dosya icerigi:\n```\n{attached_text[:20000]}\n```"
            )

        # Arama sorgusu: dosya iceriginin tamami degil, kullanicinin sorusu.
        context = self.retrieve(message if message.strip() else (attached_text or "")[:500])

        self.history.append({"role": "user", "content": user_content})
        self._trim()

        messages = [
            {"role": "system", "content": SYSTEM_TEMPLATE.format(context=context)}
        ] + self.history

        try:
            resp = self._client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.3,
                max_tokens=4096,
            )
            answer = resp.choices[0].message.content
        except Exception as exc:
            self.history.pop()  # basarisiz turu gecmiste birakma
            return f"**Model cagrisi basarisiz:** `{type(exc).__name__}: {exc}`"

        self.history.append({"role": "assistant", "content": answer})
        return answer

    def fix_error(self, error_text: str) -> str:
        """Aldigin hatayi son konusmanin devami olarak gonderir."""
        return self.ask(
            f"Onceki kodu calistirdigimda su hatayi aldim. Sebebini acikla "
            f"ve duzeltilmis tam kodu ver:\n\n```\n{error_text}\n```"
        )


# --------------------------------------------------------------------------
# Komut satirindan indeksleme:  python zafer_core.py "C:\...\all_materials"
# --------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    target = sys.argv[1] if len(sys.argv) > 1 else os.getenv("ZAFER_MATERIALS", ".")
    force = "--force" in sys.argv
    index_folder(target, force=force)